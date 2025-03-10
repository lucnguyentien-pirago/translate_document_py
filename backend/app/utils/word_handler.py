import docx
import io
from typing import List, Dict

class WordHandler:
    def __init__(self):
        pass
    
    def extract_text_from_docx(self, file_content: bytes) -> List[Dict[str, str]]:
        """
        Trích xuất văn bản từ file Word
        Trả về danh sách các đoạn, mỗi đoạn là một dictionary với key là số đoạn và value là nội dung
        """
        try:
            doc = docx.Document(io.BytesIO(file_content))
            result = []
            
            # Xử lý từng đoạn văn
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text and not paragraph.text.isspace():
                    result.append({
                        "paragraph": i + 1,
                        "content": paragraph.text
                    })
            
            return result
        except Exception as e:
            raise Exception(f"Lỗi khi xử lý file Word: {str(e)}")
    
    def create_translated_docx(self, original_file: bytes, translated_paragraphs: List[Dict[str, str]]) -> bytes:
        """
        Tạo file Word với nội dung đã được dịch
        """
        try:
            # Đọc file Word gốc
            doc = docx.Document(io.BytesIO(original_file))
            
            # Tạo mapping từ paragraph index đến translated content
            translations = {item["paragraph"]: item.get("translated_content", "") 
                          for item in translated_paragraphs 
                          if "translated_content" in item}
            
            # Cập nhật nội dung đã dịch
            for i, paragraph in enumerate(doc.paragraphs):
                if i + 1 in translations and translations[i + 1]:
                    # Giữ lại định dạng văn bản gốc
                    paragraph_format = paragraph.paragraph_format
                    runs_formats = [(run.bold, run.italic, run.underline, run.font.size) for run in paragraph.runs]
                    
                    # Cập nhật nội dung
                    paragraph.clear()
                    paragraph.add_run(translations[i + 1])
                    
                    # Áp dụng lại định dạng
                    if paragraph.runs and runs_formats:
                        for j, formats in enumerate(runs_formats):
                            if j < len(paragraph.runs):
                                run = paragraph.runs[j]
                                run.bold, run.italic, run.underline, run.font.size = formats
            
            # Lưu tài liệu vào buffer
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output.getvalue()
        except Exception as e:
            raise Exception(f"Lỗi khi tạo file Word đã dịch: {str(e)}") 